"""
Document Extraction Service — extracts text from file attachments
for compliance scanning through the CIL pipeline.

Supported formats:
  - PDF (via PyMuPDF / fitz)
  - DOCX (via python-docx)
  - CSV / TSV (via csv module)
  - Plain text (.txt, .md, .json, .xml, .yaml, .log, .py, .js, etc.)
  - Images with text (via pytesseract OCR, optional)
"""

import csv
import io
import logging
import mimetypes
from pathlib import Path

logger = logging.getLogger("aigp.document_extractor")

# Max text to extract per file (prevent memory issues)
MAX_EXTRACT_CHARS = 50_000


class DocumentExtractor:
    """Extracts readable text from various file formats."""

    def extract(self, content: bytes, filename: str, content_type: str = "") -> dict:
        """
        Extract text from file content.

        Returns:
            {
                "text": str,           # extracted text
                "filename": str,
                "content_type": str,
                "pages": int,          # for PDFs
                "method": str,         # extraction method used
                "truncated": bool,
            }
        """
        if not content_type:
            content_type, _ = mimetypes.guess_type(filename) or ("application/octet-stream", None)

        ext = Path(filename).suffix.lower()
        text = ""
        pages = 0
        method = "unknown"

        try:
            if content_type == "application/pdf" or ext == ".pdf":
                text, pages = self._extract_pdf(content)
                method = "pymupdf"

            elif content_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ) or ext in (".docx", ".doc"):
                text = self._extract_docx(content)
                method = "python-docx"

            elif content_type == "text/csv" or ext in (".csv", ".tsv"):
                text = self._extract_csv(content, delimiter="\t" if ext == ".tsv" else ",")
                method = "csv"

            elif content_type in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ) or ext in (".xlsx", ".xls"):
                text = self._extract_xlsx(content)
                method = "openpyxl"

            elif self._is_text_file(content_type, ext):
                text = content.decode("utf-8", errors="replace")
                method = "plaintext"

            elif content_type and content_type.startswith("image/"):
                text = self._extract_image_ocr(content)
                method = "tesseract-ocr"

            else:
                # Try as plain text as fallback
                try:
                    text = content.decode("utf-8", errors="strict")
                    method = "plaintext-fallback"
                except UnicodeDecodeError:
                    text = ""
                    method = "unsupported"

        except Exception as e:
            logger.warning("Extraction failed for %s (%s): %s", filename, content_type, e)
            text = ""
            method = f"error: {str(e)[:100]}"

        truncated = len(text) > MAX_EXTRACT_CHARS
        if truncated:
            text = text[:MAX_EXTRACT_CHARS]

        return {
            "text": text,
            "filename": filename,
            "content_type": content_type,
            "pages": pages,
            "method": method,
            "truncated": truncated,
            "char_count": len(text),
        }

    def _extract_pdf(self, content: bytes) -> tuple[str, int]:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed, cannot extract PDF")
            return "", 0

        doc = fitz.open(stream=content, filetype="pdf")
        pages = len(doc)
        texts = []
        for page in doc:
            texts.append(page.get_text())
        doc.close()
        return "\n".join(texts), pages

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX")
            return ""

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    def _extract_csv(self, content: bytes, delimiter: str = ",") -> str:
        """Extract text from CSV/TSV."""
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = []
        for i, row in enumerate(reader):
            if i > 500:  # Limit rows
                break
            rows.append(" | ".join(row))
        return "\n".join(rows)

    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from Excel files."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl not installed, cannot extract XLSX")
            return ""

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        texts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i > 500:
                    break
                cells = [str(c) for c in row if c is not None]
                if cells:
                    texts.append(" | ".join(cells))
        wb.close()
        return "\n".join(texts)

    def _extract_image_ocr(self, content: bytes) -> str:
        """Extract text from images using Tesseract OCR."""
        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            logger.warning("Pillow/pytesseract not installed, cannot OCR images")
            return ""

        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img)
        return text.strip()

    def _is_text_file(self, content_type: str, ext: str) -> bool:
        """Check if file is a text-based format."""
        text_types = {"text/plain", "text/html", "text/xml", "text/markdown",
                      "application/json", "application/xml", "application/yaml",
                      "application/x-yaml", "text/yaml", "text/x-python",
                      "application/javascript", "text/javascript"}
        text_exts = {".txt", ".md", ".json", ".xml", ".yaml", ".yml", ".log",
                     ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go",
                     ".rs", ".rb", ".sh", ".bat", ".ps1", ".sql", ".html", ".css",
                     ".env", ".ini", ".cfg", ".conf", ".toml"}
        return content_type in text_types or ext in text_exts
